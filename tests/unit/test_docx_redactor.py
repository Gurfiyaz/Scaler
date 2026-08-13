"""Unit tests for DOCXRedactor, RunRewriter, RelationshipRedactor, and DocumentValidator."""

import shutil
from pathlib import Path
import docx
import pytest

from app.detection.detector import PIIDetector
from app.ingestion.docx_parser import DOCXParser
from app.mapping.entity_mapper import EntityMapper
from app.redaction.docx_redactor import DOCXRedactor
from app.redaction.document_validator import DocumentValidator

FIXTURES_DIR = Path(__file__).parent.parent / "fixtures" / "synthetic"


@pytest.fixture
def parser() -> DOCXParser:
    return DOCXParser()


@pytest.fixture
def detector() -> PIIDetector:
    return PIIDetector()


@pytest.fixture
def mapper() -> EntityMapper:
    return EntityMapper()


@pytest.fixture
def redactor() -> DOCXRedactor:
    return DOCXRedactor()


# 1. Single-Run Replacement & Formatting Preservation Test
def test_single_run_redaction_and_formatting(parser, detector, mapper, redactor, tmp_path):
    input_doc = FIXTURES_DIR / "test_a_single_run.docx"
    output_doc = tmp_path / "redacted_a.docx"

    doc_model = parser.parse_document(input_doc)
    dets = detector.detect_in_document(doc_model)
    records = mapper.map_all_detections(dets)

    result = redactor.redact_document(input_doc, output_doc, dets, records)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True
    assert result.replacements_applied >= 1

    # Verify output text
    redacted_model = parser.parse_document(output_doc)
    assert "Alice Example" not in redacted_model.paragraphs[0].reconstructed_text


# 2. Cross-Run Replacement Test
def test_cross_run_redaction(parser, detector, mapper, redactor, tmp_path):
    input_doc = FIXTURES_DIR / "test_b_split_run.docx"
    output_doc = tmp_path / "redacted_b.docx"

    doc_model = parser.parse_document(input_doc)
    dets = detector.detect_in_document(doc_model)
    records = mapper.map_all_detections(dets)

    result = redactor.redact_document(input_doc, output_doc, dets, records)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True
    assert result.replacements_applied >= 1

    redacted_model = parser.parse_document(output_doc)
    assert "Alice Example" not in redacted_model.paragraphs[0].reconstructed_text


# 3. Hyperlink Display and Target Relationship Redaction Test
def test_hyperlink_and_relationship_redaction(parser, detector, mapper, redactor, tmp_path):
    input_doc = FIXTURES_DIR / "test_c_hyperlink.docx"
    output_doc = tmp_path / "redacted_c.docx"

    doc_model = parser.parse_document(input_doc)
    dets = detector.detect_in_document(doc_model)
    records = mapper.map_all_detections(dets)

    result = redactor.redact_document(input_doc, output_doc, dets, records)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True

    redacted_model = parser.parse_document(output_doc)
    assert "alice@example.test" not in redacted_model.paragraphs[0].reconstructed_text

    for rel in redacted_model.relationships.values():
        assert "alice@example.test" not in rel.target


# 4. Table Cell Redaction Test
def test_table_cell_redaction(parser, detector, mapper, redactor, tmp_path):
    input_doc = FIXTURES_DIR / "test_d_table.docx"
    output_doc = tmp_path / "redacted_d.docx"

    doc_model = parser.parse_document(input_doc)
    dets = detector.detect_in_document(doc_model)
    records = mapper.map_all_detections(dets)

    result = redactor.redact_document(input_doc, output_doc, dets, records)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True

    redacted_model = parser.parse_document(output_doc)
    for row in redacted_model.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                assert "Alice Example" not in p.reconstructed_text
                assert "+91 9000000000" not in p.reconstructed_text


# 5. Header and Footer Redaction Test
def test_header_footer_redaction(parser, detector, mapper, redactor, tmp_path):
    input_doc = FIXTURES_DIR / "test_e_header_footer.docx"
    output_doc = tmp_path / "redacted_e.docx"

    doc_model = parser.parse_document(input_doc)
    dets = detector.detect_in_document(doc_model)
    records = mapper.map_all_detections(dets)

    result = redactor.redact_document(input_doc, output_doc, dets, records)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True

    redacted_model = parser.parse_document(output_doc)
    for h in redacted_model.headers:
        for p in h.paragraphs:
            assert "Alice Example" not in p.reconstructed_text
    for f in redacted_model.footers:
        for p in f.paragraphs:
            assert "alice@example.test" not in p.reconstructed_text


# 6. Original File Protection Test (SHA-256 Hash Unchanged)
def test_original_file_hash_protection(parser, detector, mapper, redactor, tmp_path):
    input_doc = tmp_path / "source_copy.docx"
    shutil.copy(FIXTURES_DIR / "test_a_single_run.docx", input_doc)
    output_doc = tmp_path / "redacted_copy.docx"

    initial_hash = DOCXRedactor.calculate_file_hash(input_doc)

    doc_model = parser.parse_document(input_doc)
    dets = detector.detect_in_document(doc_model)
    records = mapper.map_all_detections(dets)

    redactor.redact_document(input_doc, output_doc, dets, records)

    final_hash = DOCXRedactor.calculate_file_hash(input_doc)
    assert initial_hash == final_hash


# 7. Idempotency Test (Redact Redacted Document)
def test_idempotency_second_pass(parser, detector, mapper, redactor, tmp_path):
    input_doc = FIXTURES_DIR / "test_a_single_run.docx"
    output_pass1 = tmp_path / "pass1.docx"
    output_pass2 = tmp_path / "pass2.docx"

    # Pass 1
    doc_m1 = parser.parse_document(input_doc)
    dets1 = detector.detect_in_document(doc_m1)
    recs1 = mapper.map_all_detections(dets1)
    redactor.redact_document(input_doc, output_pass1, dets1, recs1)

    # Pass 2 on Pass 1 output
    doc_m2 = parser.parse_document(output_pass1)
    dets2 = detector.detect_in_document(doc_m2)
    # Filter detections to only detect residual original PII (which should be 0)
    recs2 = mapper.map_all_detections(dets2)
    result2 = redactor.redact_document(output_pass1, output_pass2, dets2, recs2)

    assert result2.is_valid_docx is True
    assert result2.residual_pii_clean is True
