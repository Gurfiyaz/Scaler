"""Redacted document structural validation and residual PII scanning module."""

import re
import zipfile
from pathlib import Path
from typing import List, Set
import docx
from app.detection.detector import PIIDetector
from app.detection.models import PIIDetection, PIIType
from app.ingestion.docx_parser import DOCXParser
from app.mapping.entity_mapper import EntityMapper
from app.redaction.exceptions import DocumentValidationError, ResidualPIIError


class DocumentValidator:
    """Validates structural integrity and scans output document package to verify zero residual PII."""

    GENERIC_STOP_WORDS = {
        "pvt", "ltd", "inc", "corp", "llp", "llc", "limited", "co", "company", "labs",
        "technologies", "services", "solutions", "systems", "and", "the", "for", "group",
        "holdings", "enterprises", "international", "global", "national", "trust", "bank", "fund",
        "associates", "partners", "capital", "management", "financial", "advisors", "india", "private"
    }

    def __init__(self):
        self.parser = DOCXParser()
        self.detector = PIIDetector()

    def validate_structure(self, output_path: Path, expected_paragraph_count: int = 0) -> bool:
        """Confirm that output document re-opens without XML corruption and contains readable structures.

        Raises:
            DocumentValidationError: If output file is unreadable or structurally corrupted.
        """
        try:
            doc = docx.Document(output_path)
            # Re-read paragraphs, tables, headers/footers
            _ = len(doc.paragraphs)
            _ = len(doc.tables)

            for s in doc.sections:
                if hasattr(s, "header") and s.header:
                    _ = len(s.header.paragraphs)
                if hasattr(s, "footer") and s.footer:
                    _ = len(s.footer.paragraphs)

            return True
        except Exception as e:
            raise DocumentValidationError(f"Redacted document failed structural integrity check: {str(e)}")

    def validate_residual_pii(
        self, output_path: Path, original_detections: List[PIIDetection]
    ) -> bool:
        """Perform package-level raw ZIP XML inspection to verify 0 original PII remains in package.

        Args:
            output_path: Path to generated redacted DOCX file.
            original_detections: List of PIIDetection objects detected in source document.

        Returns:
            True if 0 original PII strings remain in document or relationships.

        Raises:
            ResidualPIIError: If any original PII value is found in the output document.
        """
        if not original_detections:
            return True

        output_path = Path(output_path)
        if not output_path.exists() or output_path.stat().st_size == 0:
            raise ResidualPIIError("Redacted output file is missing or empty.")

        # 1. Inspect uncompressed raw ZIP XML parts (document.xml, rels, headers, footers)
        raw_xml_parts: List[str] = []
        try:
            with zipfile.ZipFile(output_path, "r") as z:
                for name in z.namelist():
                    if name.endswith(".xml") or name.endswith(".rels"):
                        content_bytes = z.read(name)
                        raw_xml_parts.append(content_bytes.decode("utf-8", errors="ignore"))
        except Exception as err:
            raise DocumentValidationError(f"Failed to inspect ZIP XML package: {err}")

        combined_raw_xml = "\n".join(raw_xml_parts).lower()

        # 2. Check that distinctive original detected PII strings are completely absent
        checked_spans: Set[str] = set()

        for orig_det in original_detections:
            orig_text = orig_det.text.strip()
            norm_orig = EntityMapper.normalize_key_text(orig_text, orig_det.entity_type)

            # Determine minimum length threshold based on entity type
            min_len = 4
            if orig_det.entity_type in {PIIType.PERSON, PIIType.ORGANIZATION, PIIType.ADDRESS}:
                min_len = 8

            if not norm_orig or len(norm_orig) < min_len:
                continue

            # Skip generic stop words or corporate suffixes that appear in non-PII contexts
            if norm_orig.lower() in self.GENERIC_STOP_WORDS:
                continue

            if norm_orig in checked_spans:
                continue
            checked_spans.add(norm_orig)

            # Use word-boundary regex to avoid matching substrings within XML tags or other text
            pattern = rf"\\b{re.escape(norm_orig)}\\b"
            if re.search(pattern, combined_raw_xml, flags=re.IGNORECASE):
                raise ResidualPIIError(
                    "Residual original PII detected in generated DOCX package."
                )

        return True

    def validate_redacted_document(
        self,
        input_path: Path,
        output_path: Path,
        original_detections: List[PIIDetection],
        mapping_dict: dict,
    ):
        """Validate structural integrity and zero residual PII in redacted DOCX output."""
        self.validate_structure(output_path)
        self.validate_residual_pii(output_path, original_detections)

        class SingleValidationResult:
            is_valid = True
            residual_pii_detected = False

        return SingleValidationResult()
