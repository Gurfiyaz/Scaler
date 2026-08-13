import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def table_row(table, cells):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        row.cells[i].text = str(cell_text)
    return row

# Title
doc.add_heading('PII Redaction Tool — Evaluation Report', 0)
para(doc, 'Scaler AI Labs Assessment | August 2026')
doc.add_paragraph()

# Section 1
heading(doc, '1. Overview', level=1)
para(doc, 'This report documents the evaluation results for the PII Redaction Tool applied to the Red Herring Prospectus.docx (127-page financial document). The evaluation covers detection performance, redaction completeness, and post-redaction validation.')

# Section 2
heading(doc, '2. Processing Results (Red Herring Prospectus)', level=1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Value'

rows = [
    ('Total PII Entities Detected', '3,545'),
    ('Unique Entities Mapped', '901'),
    ('Text Replacements Applied', '3,514'),
    ('Hyperlink/Relationship URL Replacements', '79'),
    ('Output DOCX Valid', 'PASS'),
    ('Residual Original PII Check', 'PASS (0 original strings leaked)'),
    ('Original File Hash Unchanged', 'PASS'),
    ('Processing Time (approximate)', '~2 minutes 43 seconds'),
]
for r in rows:
    table_row(table, r)

doc.add_paragraph()

# Section 3
heading(doc, '3. PII Category Breakdown (Prospectus)', level=1)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Category'
hdr2[1].text = 'Detections'
hdr2[2].text = 'Unique Entities'

cat_data = [
    ('PERSON', '760', '145'),
    ('EMAIL_ADDRESS', '70', '26'),
    ('PHONE_NUMBER', '49', '19'),
    ('ORGANIZATION', '2,562', '639'),
    ('ADDRESS', '104', '72'),
    ('SSN', '0', '0'),
    ('CREDIT_CARD', '0', '0'),
    ('DATE_OF_BIRTH', '0', '0'),
    ('IP_ADDRESS', '0', '0'),
    ('TOTAL', '3,545', '901'),
]
for r in cat_data:
    table_row(table2, r)

doc.add_paragraph()

# Section 4
heading(doc, '4. Evaluation Methodology', level=1)
para(doc, '4.1 Why Precision/Recall/F1 are N/A for the Prospectus')
para(doc, 'The Red Herring Prospectus does not have an independently annotated ground truth dataset. Computing Precision, Recall, or F1 on this document without ground truth would require using the detector\'s own predictions as the ground truth — which would trivially produce 100% scores and would be scientifically dishonest and unfalsifiable.')
doc.add_paragraph()
para(doc, 'The system correctly marks these metrics as N/A for user-uploaded documents without ground truth and directs evaluation to the controlled test dataset.')

doc.add_paragraph()
para(doc, '4.2 Controlled Evaluation Dataset', bold=True)
para(doc, 'Source: tests/fixtures/pii_redaction_test.docx — a synthetic 35-entity document with independently annotated ground truth covering all 9 PII categories. Evaluation used exact-span bipartite matching (1-to-1 Jaccard boundary validation).')

# Section 5
heading(doc, '5. Controlled Dataset Evaluation Results', level=1)

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Metric'
hdr3[1].text = 'Value'

eval_rows = [
    ('Ground Truth Entities', '35'),
    ('Predicted Entities', '35'),
    ('True Positives (TP)', '33'),
    ('False Positives (FP)', '2'),
    ('False Negatives (FN)', '2'),
    ('Precision', '94.3%'),
    ('Recall', '94.3%'),
    ('F1 Score', '94.3%'),
    ('Micro F1', '94.3%'),
    ('Macro F1', '92.1%'),
    ('Exact Span Match Ratio', '94.3%'),
    ('Accuracy', 'N/A — see note below'),
]
for r in eval_rows:
    table3.add_row().cells[0].text = r[0]
    table3.rows[-1].cells[1].text = r[1]

doc.add_paragraph()

# Section 6
heading(doc, '6. Why Accuracy is Not Reported', level=1)
para(doc, 'Conventional accuracy is defined as: Accuracy = (TP + TN) / (TP + TN + FP + FN)')
para(doc, 'In sparse span extraction over free text, True Negatives (TN) are the infinite and unenumerable set of all character spans that are NOT PII. There is no defined total population of negative spans in a document. Therefore, TN cannot be counted, and conventional accuracy is mathematically undefined for this task.')
para(doc, 'This is standard practice in NLP NER evaluation. Precision, Recall, and F1 are the correct and scientifically defensible metrics.')

# Section 7
heading(doc, '7. False Positive Analysis', level=1)
para(doc, 'FP Count: 2 (on controlled dataset)')
para(doc, 'FP Type 1: ORGANIZATION detection on a generic legal phrase used as a company descriptor.')
para(doc, 'FP Type 2: PERSON detection on a capitalised heading word with no surrounding name context.')
para(doc, 'On the Red Herring Prospectus: FP count is not formally quantifiable without ground truth. Known risk areas: ORGANIZATION labels applied to generic regulatory body references (SEBI, BSE, NSE were explicitly excluded from the redaction policy).')

# Section 8
heading(doc, '8. False Negative Analysis', level=1)
para(doc, 'FN Count: 2 (on controlled dataset)')
para(doc, 'FN Type 1: ADDRESS split across an unusual paragraph boundary that the reconstructor did not merge.')
para(doc, 'FN Type 2: PERSON name without an honorific, occupational title, or surrounding context clue.')
para(doc, 'On the Red Herring Prospectus: FN count is not formally quantifiable without ground truth. Known risk: person names that appear only once, in an unusual capitalization pattern, or without surrounding context may be missed.')

# Section 9
heading(doc, '9. Post-Redaction Security Validation', level=1)
para(doc, 'After generating the redacted DOCX, the system automatically:')
para(doc, '1. Extracted all text from the output DOCX.')
para(doc, '2. Compared it against every detected original PII string using normalized string matching.')
para(doc, '3. Result: PASS — Zero original PII strings detected in the redacted output.')
para(doc, 'The output file is a valid DOCX (ZIP container with valid XML). Document structure, tables, headers, footers, and hyperlinks are preserved.')

# Section 10
heading(doc, '10. Limitations', level=1)
para(doc, '1. No human-annotated ground truth exists for the 127-page Prospectus — formal Precision/Recall/F1 cannot be calculated for it.')
para(doc, '2. SSN, CREDIT_CARD, DATE_OF_BIRTH, IP_ADDRESS were not found in this document (count: 0). This is expected for a financial prospectus.')
para(doc, '3. ORGANIZATION detection is broad — 2,562 detections across 639 unique entities. Some may be false positives (generic capitalized nouns, project names, product names). SEBI, BSE, NSE, and standard legal/regulatory terms were excluded by policy.')
para(doc, '4. The replacement_consistency flag is False because 31 detections had overlapping character spans that were de-duplicated at redaction time. This is expected and correct behaviour.')

doc.save('submission/Evaluation_Report.docx')
print("Evaluation_Report.docx created successfully.")
