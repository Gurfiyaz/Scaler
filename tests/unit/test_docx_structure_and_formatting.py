"""DOCX Structure, Formatting, Media & Unicode Edge Case Unit Tests."""

from pathlib import Path
import docx
import pytest
from app.detection.detector import PIIDetector
from app.detection.models import PIIType
from app.ingestion.docx_parser import DOCXParser
from app.mapping.entity_mapper import EntityMapper
from app.redaction.docx_redactor import DOCXRedactor

FIXTURES_DIR = Path(__file__).parent.parent / "fixtures" / "synthetic"


@pytest.fixture
def parser():
    return DOCXParser()


@pytest.fixture
def detector():
    return PIIDetector()


@pytest.fixture
def mapper():
    return EntityMapper()


@pytest.fixture
def redactor():
    return DOCXRedactor()


# 1. Unicode & Special Character Processing Test
def test_unicode_and_special_character_processing(parser, detector, mapper, redactor, tmp_path):
    text = "Contact Francois Rene or Maria-Jose at francois.rene@example.com."
    doc = docx.Document()
    doc.add_paragraph(text)

    in_doc = tmp_path / "unicode_test.docx"
    out_doc = tmp_path / "redacted_unicode.docx"
    doc.save(in_doc)

    doc_model = parser.parse_document(in_doc)
    dets = detector.detect_in_document(doc_model)
    recs = mapper.map_all_detections(dets)

    result = redactor.redact_document(in_doc, out_doc, dets, recs)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True

    redacted_model = parser.parse_document(out_doc)
    redacted_text = redacted_model.paragraphs[0].reconstructed_text
    assert "francois.rene@example.com" not in redacted_text


# 2. Synthetic Formatting Preservation Verification
def test_formatting_preservation_fixture(parser, detector, mapper, redactor, tmp_path):
    in_doc = FIXTURES_DIR / "test_f_formatting.docx"
    out_doc = tmp_path / "redacted_f.docx"

    doc_model = parser.parse_document(in_doc)
    dets = detector.detect_in_document(doc_model)
    recs = mapper.map_all_detections(dets)

    result = redactor.redact_document(in_doc, out_doc, dets, recs)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True

    redacted_docx = docx.Document(out_doc)
    p = redacted_docx.paragraphs[0]
    assert len(p.runs) >= 1
    assert "Alice Example" not in p.text


# 3. Dynamic Hyperlink and Relationship Redaction Verification
def test_hyperlink_dynamic_relationship_redaction(parser, detector, mapper, redactor, tmp_path):
    in_doc = FIXTURES_DIR / "test_c_hyperlink.docx"
    out_doc = tmp_path / "redacted_c.docx"

    doc_model = parser.parse_document(in_doc)
    dets = detector.detect_in_document(doc_model)
    recs = mapper.map_all_detections(dets)

    result = redactor.redact_document(in_doc, out_doc, dets, recs)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True

    redacted_model = parser.parse_document(out_doc)
    for r in redacted_model.relationships.values():
        assert "alice@example.test" not in r.target


# 4. Multi-Row Table Cell Redaction Test
def test_table_multi_cell_redaction(parser, detector, mapper, redactor, tmp_path):
    in_doc = FIXTURES_DIR / "test_d_table.docx"
    out_doc = tmp_path / "redacted_d.docx"

    doc_model = parser.parse_document(in_doc)
    dets = detector.detect_in_document(doc_model)
    recs = mapper.map_all_detections(dets)

    result = redactor.redact_document(in_doc, out_doc, dets, recs)

    assert result.is_valid_docx is True
    assert result.residual_pii_clean is True

    redacted_model = parser.parse_document(out_doc)
    for t in redacted_model.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    assert "Alice Example" not in p.reconstructed_text
                    assert "+91 9000000000" not in p.reconstructed_text
