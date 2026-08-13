"""Generator for synthetic DOCX test fixtures.

PRIVACY RULE: Contains ONLY fake/synthetic placeholder data.
No real PII is ever used in test fixtures.
"""

from pathlib import Path
import docx
from docx.shared import Pt, RGBColor


def create_synthetic_fixtures(output_dir: Path):
    """Generate 9 synthetic test files for integration testing."""
    output_dir.mkdir(parents=True, exist_ok=True)

    # Fixture 1: Single Run PII
    doc_a = docx.Document()
    doc_a.add_paragraph("Contact Alice Example today for details.")
    doc_a.save(output_dir / "test_a_single_run.docx")

    # Fixture 2: Split Run PERSON
    doc_b = docx.Document()
    p_b = doc_b.add_paragraph()
    p_b.add_run("Alice ")
    p_b.add_run("Example")
    doc_b.save(output_dir / "test_b_split_run.docx")

    # Fixture 3: Email Hyperlink with Relationship Target
    doc_c = docx.Document()
    p_c = doc_c.add_paragraph("Email us at ")
    p_elem = p_c._p
    part = doc_c.part
    r_id = part.relate_to(
        "mailto:alice@example.test",
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.ns.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    color = docx.oxml.shared.OxmlElement('w:color')
    color.set(docx.oxml.ns.qn('w:val'), '0000FF')
    rPr.append(color)
    new_run.append(rPr)
    text_elem = docx.oxml.shared.OxmlElement('w:t')
    text_elem.text = "alice@example.test"
    new_run.append(text_elem)
    hyperlink.append(new_run)
    p_elem.append(hyperlink)
    doc_c.save(output_dir / "test_c_hyperlink.docx")

    # Fixture 4: Table PII
    doc_d = docx.Document()
    tbl = doc_d.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Employee Name"
    tbl.cell(0, 1).text = "Contact Phone"
    tbl.cell(1, 0).text = "Alice Example"
    tbl.cell(1, 1).text = "+91 9000000000"
    doc_d.save(output_dir / "test_d_table.docx")

    # Fixture 5: Header / Footer PII
    doc_e = docx.Document()
    sec = doc_e.sections[0]
    sec.header.paragraphs[0].text = "Header Confidential - Alice Example"
    sec.footer.paragraphs[0].text = "Footer Contact: alice@example.test"
    doc_e.add_paragraph("Body content of header footer test.")
    doc_e.save(output_dir / "test_e_header_footer.docx")

    # Fixture 6: Formatting Preservation
    doc_f = docx.Document()
    p_f = doc_f.add_paragraph()
    r_bold = p_f.add_run("Alice ")
    r_bold.bold = True
    r_italic = p_f.add_run("Example ")
    r_italic.italic = True
    r_underline = p_f.add_run("is contacting us.")
    r_underline.underline = True
    r_color = p_f.add_run(" Colored text.")
    r_color.font.color.rgb = RGBColor(255, 0, 0)
    doc_f.save(output_dir / "test_f_formatting.docx")

    # Fixture 7: Multiple Spans in One Paragraph
    doc_g = docx.Document()
    doc_g.add_paragraph("Contact Alice Example at alice@example.test or call +91 9000000000.")
    doc_g.save(output_dir / "test_g_multiple_spans.docx")

    # Fixture 8: Multiple Occurrences of Same Entity
    doc_h = docx.Document()
    doc_h.add_paragraph("Alice Example appeared here and Alice Example appeared there.")
    doc_h.save(output_dir / "test_h_multiple_occurrences.docx")

    # Fixture 9: Adjacent PII Spans
    doc_i = docx.Document()
    doc_i.add_paragraph("Alice Example alice@example.test")
    doc_i.save(output_dir / "test_i_adjacent_spans.docx")


if __name__ == "__main__":
    fixtures_path = Path(__file__).parent
    create_synthetic_fixtures(fixtures_path)
    print(f"Generated 9 synthetic fixtures in '{fixtures_path}'")
