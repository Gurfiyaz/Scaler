"""Master DOCX Redaction Engine orchestrating XML run rewriting, relationship targets, and validation."""

import hashlib
import time
from pathlib import Path
from typing import Dict, List, Optional
import docx
from app.core.logging_config import logger
from app.detection.models import PIIDetection
from app.ingestion.models import SourceLocationType
from app.mapping.entity_mapper import EntityMapper
from app.mapping.models import EntityMappingKey, EntityMappingRecord
from app.redaction.document_validator import DocumentValidator
from app.redaction.exceptions import RedactionError
from app.redaction.hyperlink_redactor import HyperlinkRedactor
from app.redaction.models import RedactionResult, RedactionTask
from app.redaction.run_rewriter import RunRewriter


class DOCXRedactor:
    """Orchestrates XML run rewriting, relationship target redactions, and output integrity validation."""

    def __init__(self):
        self.validator = DocumentValidator()

    @classmethod
    def calculate_file_hash(cls, file_path: Path) -> str:
        """Compute SHA-256 hash of a file for integrity verification."""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            while chunk := f.read(65536):
                sha256.update(chunk)
        return sha256.hexdigest()

    def redact_document(
        self,
        input_path: Path,
        output_path: Path,
        detections: List[PIIDetection],
        mapping_records: Dict[EntityMappingKey, EntityMappingRecord],
    ) -> RedactionResult:
        """Execute end-to-end redaction on input DOCX file and write to output_path.

        CRITICAL SAFETY RULE: Never modifies input_path in place. Verifies input SHA-256 hash before and after.

        Args:
            input_path: Path to source DOCX file.
            output_path: Path to destination redacted DOCX file.
            detections: List of PIIDetection objects from Phase 4.
            mapping_records: Dictionary of EntityMappingRecord objects from Phase 5.

        Returns:
            RedactionResult summary.
        """
        input_path = Path(input_path).resolve()
        output_path = Path(output_path).resolve()

        if input_path == output_path:
            raise RedactionError("Output path must not be identical to input path.")

        if not input_path.exists():
            raise RedactionError(f"Input document not found: {input_path}")

        # Compute initial input file hash
        initial_hash = self.calculate_file_hash(input_path)

        t_start = time.perf_counter()

        # Open python-docx Document
        doc = docx.Document(input_path)

        # Build RedactionTask objects grouped by paragraph location
        tasks_by_para: Dict[tuple[SourceLocationType, int], List[RedactionTask]] = {}

        for det in detections:
            norm_text = EntityMapper.normalize_key_text(det.text, det.entity_type)
            key = EntityMappingKey(entity_type=det.entity_type, normalized_original=norm_text)
            if key not in mapping_records:
                continue

            rec = mapping_records[key]
            task = RedactionTask(
                paragraph_index=det.paragraph_index,
                start_offset=det.start,
                end_offset=det.end,
                original_text=det.text,
                replacement_text=rec.replacement_value,
                entity_type=det.entity_type,
                location=det.location,
                run_indices=det.run_indices,
            )

            loc_key = (det.location.location_type, det.location.paragraph_index)
            tasks_by_para.setdefault(loc_key, []).append(task)

        total_applied = 0
        para_counter = 0

        # 1. Redact Body Paragraphs
        for doc_p in doc.paragraphs:
            loc_key = (SourceLocationType.BODY, para_counter)
            if loc_key in tasks_by_para:
                applied = RunRewriter.apply_paragraph_redactions(doc_p, tasks_by_para[loc_key])
                total_applied += applied
            para_counter += 1

        # 2. Redact Table Cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for doc_p in cell.paragraphs:
                        loc_key = (SourceLocationType.TABLE, para_counter)
                        if loc_key in tasks_by_para:
                            applied = RunRewriter.apply_paragraph_redactions(doc_p, tasks_by_para[loc_key])
                            total_applied += applied
                        para_counter += 1

        # 3. Redact Headers and Footers
        for sec in doc.sections:
            for header_attr in ["header", "first_page_header", "even_page_header"]:
                if hasattr(sec, header_attr):
                    hdr = getattr(sec, header_attr)
                    if hdr and hdr.paragraphs:
                        for doc_p in hdr.paragraphs:
                            loc_key = (SourceLocationType.HEADER, para_counter)
                            if loc_key in tasks_by_para:
                                applied = RunRewriter.apply_paragraph_redactions(doc_p, tasks_by_para[loc_key])
                                total_applied += applied
                            para_counter += 1

            for footer_attr in ["footer", "first_page_footer", "even_page_footer"]:
                if hasattr(sec, footer_attr):
                    ftr = getattr(sec, footer_attr)
                    if ftr and ftr.paragraphs:
                        for doc_p in ftr.paragraphs:
                            loc_key = (SourceLocationType.FOOTER, para_counter)
                            if loc_key in tasks_by_para:
                                applied = RunRewriter.apply_paragraph_redactions(doc_p, tasks_by_para[loc_key])
                                total_applied += applied
                            para_counter += 1

        # 4. Redact Relationship Target URIs (.rels)
        rel_applied = HyperlinkRedactor.redact_hyperlinks_and_relationships(
            doc, list(mapping_records.values())
        )

        # Save to output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        redaction_duration = time.perf_counter() - t_start

        # 5. Output Validation (Structural & Residual PII)
        v_start = time.perf_counter()
        is_valid = self.validator.validate_structure(output_path, len(doc.paragraphs))
        residual_clean = self.validator.validate_residual_pii(output_path, detections)
        val_duration = time.perf_counter() - v_start

        # Confirm input file hash remains 100% byte-for-byte untouched
        final_hash = self.calculate_file_hash(input_path)
        if initial_hash != final_hash:
            raise RedactionError("CRITICAL AUDIT FAILURE: Input file was modified during redaction!")

        logger.info(
            f"Document Redaction Complete. Applied {total_applied} text replacements and {rel_applied} relationship updates. "
            f"Output saved to '{output_path.name}'."
        )

        return RedactionResult(
            total_detections=len(detections),
            replacements_applied=total_applied,
            output_path=str(output_path),
            is_valid_docx=is_valid,
            residual_pii_clean=residual_clean,
            timing={
                "redaction_sec": redaction_duration,
                "validation_sec": val_duration,
                "total_sec": redaction_duration + val_duration,
            },
        )
