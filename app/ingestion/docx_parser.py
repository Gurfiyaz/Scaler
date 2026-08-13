"""DOCX Document Parser and Ingestion Engine.

Provides complete read-only parsing of DOCX body paragraphs, tables, headers, footers,
run character offset reconstruction, and dynamic hyperlink relationship resolution.
"""

import io
import os
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional, Union

import docx
from docx.opc.exceptions import OpcError
from docx.oxml.exceptions import InvalidXmlError

from app.core.logging_config import logger
from app.ingestion.exceptions import (
    CorruptedDocumentError,
    DocumentNotFoundError,
    EmptyDocumentError,
    InvalidDocumentError,
)
from app.ingestion.models import (
    DocumentModel,
    FormattingMetadata,
    HeaderFooterModel,
    HyperlinkModel,
    ParagraphModel,
    RelationshipModel,
    RunModel,
    SourceLocation,
    SourceLocationType,
    TableCellModel,
    TableModel,
    TableRowModel,
)
from app.ingestion.relationship_parser import RelationshipParser


class DOCXParser:
    """Read-only ingestion engine for DOCX documents."""

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    def __init__(self):
        self._relationship_parser = RelationshipParser()

    def parse_document(
        self, source: Union[str, Path, io.BytesIO], document_id: str = "doc_001"
    ) -> DocumentModel:
        """Parse a DOCX file path or binary stream into a DocumentModel.

        Args:
            source: Absolute file path, Path object, or BytesIO stream of the DOCX document.
            document_id: Optional identifier string for this document.

        Returns:
            DocumentModel containing reconstructed paragraphs, runs, tables, headers,
            footers, hyperlinks, and relationships.

        Raises:
            DocumentNotFoundError: If the file path does not exist.
            InvalidDocumentError: If the file is not a valid DOCX container.
            CorruptedDocumentError: If XML structures or ZIP contents are damaged.
            EmptyDocumentError: If document contains no text content.
        """
        file_name: Optional[str] = None

        if isinstance(source, (str, Path)):
            path_obj = Path(source)
            if not path_obj.exists():
                raise DocumentNotFoundError(f"Document file not found: '{source}'")
            file_name = path_obj.name
            try:
                with open(path_obj, "rb") as f:
                    content_bytes = f.read()
            except Exception as err:
                raise CorruptedDocumentError(f"Failed to read file '{source}': {err}")
            stream = io.BytesIO(content_bytes)
        elif isinstance(source, io.BytesIO):
            stream = source
            file_name = "stream.docx"
        else:
            raise InvalidDocumentError("Source must be a file path or BytesIO stream.")

        # Validate ZIP container
        if not zipfile.is_zipfile(stream):
            raise InvalidDocumentError("Input file is not a valid ZIP/DOCX archive.")

        stream.seek(0)
        try:
            with zipfile.ZipFile(stream, "r") as z:
                zip_names = z.namelist()
                if "word/document.xml" not in zip_names:
                    raise InvalidDocumentError("ZIP container is missing 'word/document.xml'.")

                # Extract relationships
                relationships = self._relationship_parser.parse_relationships_from_zip(
                    z, "word/_rels/document.xml.rels"
                )
        except zipfile.BadZipFile as err:
            raise CorruptedDocumentError(f"Corrupted ZIP archive: {err}")
        except InvalidDocumentError:
            raise
        except Exception as err:
            raise CorruptedDocumentError(f"Error inspecting DOCX structure: {err}")

        # Parse using python-docx
        stream.seek(0)
        try:
            doc = docx.Document(stream)
        except (InvalidXmlError, OpcError) as err:
            raise CorruptedDocumentError(f"Malformed DOCX XML: {err}")
        except Exception as err:
            raise InvalidDocumentError(f"python-docx failed to open document: {err}")

        # Construct internal representation
        body_paragraphs: List[ParagraphModel] = []
        tables_models: List[TableModel] = []
        headers_models: List[HeaderFooterModel] = []
        footers_models: List[HeaderFooterModel] = []
        hyperlinks: List[HyperlinkModel] = []

        para_counter = 0

        # 1. Body Paragraphs & Hyperlinks
        for doc_p in doc.paragraphs:
            loc = SourceLocation(
                location_type=SourceLocationType.BODY, paragraph_index=para_counter
            )
            p_model = self._parse_paragraph(doc_p, para_counter, loc)
            body_paragraphs.append(p_model)
            para_counter += 1

            # Extract hyperlinks in this paragraph
            p_links = self._extract_paragraph_hyperlinks(doc_p, p_model, relationships)
            hyperlinks.extend(p_links)

        # 2. Tables
        for table_idx, table in enumerate(doc.tables):
            t_model = self._parse_table(table, table_idx, para_counter, relationships, hyperlinks)
            tables_models.append(t_model)
            # Update paragraph index counter for paragraphs inside table cells
            for row in t_model.rows:
                for cell in row.cells:
                    para_counter += len(cell.paragraphs)

        # 3. Headers and Footers
        for section_idx, section in enumerate(doc.sections):
            for header_attr, kind_name in [
                ("header", "primary_header"),
                ("first_page_header", "first_header"),
                ("even_page_header", "even_header"),
            ]:
                if hasattr(section, header_attr):
                    hdr = getattr(section, header_attr)
                    if hdr and hdr.paragraphs:
                        hdr_model = self._parse_header_footer(
                            hdr.paragraphs, kind_name, para_counter, relationships, hyperlinks
                        )
                        if hdr_model.paragraphs:
                            headers_models.append(hdr_model)
                            para_counter += len(hdr_model.paragraphs)

            for footer_attr, kind_name in [
                ("footer", "primary_footer"),
                ("first_page_footer", "first_footer"),
                ("even_page_footer", "even_footer"),
            ]:
                if hasattr(section, footer_attr):
                    ftr = getattr(section, footer_attr)
                    if ftr and ftr.paragraphs:
                        ftr_model = self._parse_header_footer(
                            ftr.paragraphs, kind_name, para_counter, relationships, hyperlinks
                        )
                        if ftr_model.paragraphs:
                            footers_models.append(ftr_model)
                            para_counter += len(ftr_model.paragraphs)

        total_runs = sum(len(p.runs) for p in body_paragraphs)
        for tbl in tables_models:
            for r in tbl.rows:
                for c in r.cells:
                    total_runs += sum(len(p.runs) for p in c.paragraphs)

        if not body_paragraphs and not tables_models:
            raise EmptyDocumentError("Document contains no body paragraphs or tables.")

        return DocumentModel(
            document_id=document_id,
            file_name=file_name,
            paragraphs=body_paragraphs,
            tables=tables_models,
            headers=headers_models,
            footers=footers_models,
            hyperlinks=hyperlinks,
            relationships=relationships,
            total_paragraphs=para_counter,
            total_runs=total_runs,
        )

    def _parse_paragraph(
        self, doc_p: docx.text.paragraph.Paragraph, p_index: int, loc: SourceLocation
    ) -> ParagraphModel:
        runs_models: List[RunModel] = []
        current_offset = 0
        style_name = doc_p.style.name if doc_p.style else None

        # Collect all runs including those nested inside <w:hyperlink> nodes
        xml_runs = [docx.text.run.Run(r_elem, doc_p) for r_elem in doc_p._p.xpath(".//w:r")]

        for r_index, run in enumerate(xml_runs):
            r_text = run.text or ""
            r_len = len(r_text)
            start_off = current_offset
            end_off = current_offset + r_len
            current_offset = end_off

            # Extract formatting metadata
            fmt = self._extract_formatting(run)

            r_model = RunModel(
                run_index=r_index,
                paragraph_index=p_index,
                text=r_text,
                start_offset=start_off,
                end_offset=end_off,
                formatting=fmt,
            )
            runs_models.append(r_model)

        reconstructed_text = "".join(r.text for r in runs_models)

        return ParagraphModel(
            paragraph_index=p_index,
            reconstructed_text=reconstructed_text,
            runs=runs_models,
            style_name=style_name,
            location=loc,
        )

    def _extract_formatting(self, run: docx.text.run.Run) -> FormattingMetadata:
        """Extract font, size, color, bold, italic, underline properties from a Run."""
        bold = bool(run.bold)
        italic = bool(run.italic)
        underline = bool(run.underline)
        font_name = run.font.name if run.font else None
        font_size_pt = run.font.size.pt if (run.font and run.font.size) else None

        color_hex: Optional[str] = None
        if run.font and run.font.color and run.font.color.rgb:
            color_hex = str(run.font.color.rgb)

        style_name = run.style.name if run.style else None

        return FormattingMetadata(
            bold=bold,
            italic=italic,
            underline=underline,
            font_name=font_name,
            font_size_pt=font_size_pt,
            color_hex=color_hex,
            style_name=style_name,
        )

    def _extract_paragraph_hyperlinks(
        self,
        doc_p: docx.text.paragraph.Paragraph,
        p_model: ParagraphModel,
        relationships: Dict[str, RelationshipModel],
    ) -> List[HyperlinkModel]:
        """Extract hyperlink elements (<w:hyperlink>) from paragraph XML."""
        links: List[HyperlinkModel] = []
        p_element = doc_p._p

        # Find all w:hyperlink child tags
        hyperlink_elems = p_element.findall(f".//{{{self.W_NS}}}hyperlink")
        for h_elem in hyperlink_elems:
            r_id = h_elem.attrib.get(f"{{{self.R_NS}}}id")
            if not r_id:
                continue

            # Extract display text from child <w:r>/<w:t> nodes
            t_nodes = h_elem.findall(f".//{{{self.W_NS}}}t")
            display_text = "".join(t.text for t in t_nodes if t.text is not None)

            # Resolve relationship target
            rel_info = relationships.get(r_id)
            target_uri = rel_info.target if rel_info else ""
            target_mode = rel_info.target_mode if rel_info else "External"
            rel_type = rel_info.type if rel_info else "hyperlink"

            links.append(
                HyperlinkModel(
                    relationship_id=r_id,
                    display_text=display_text,
                    target_uri=target_uri,
                    relationship_type=rel_type,
                    target_mode=target_mode,
                    location=p_model.location,
                )
            )

        return links

    def _parse_table(
        self,
        table: docx.table.Table,
        table_idx: int,
        start_p_idx: int,
        relationships: Dict[str, RelationshipModel],
        hyperlinks: List[HyperlinkModel],
    ) -> TableModel:
        """Parse a table structure into TableModel."""
        rows_models: List[TableRowModel] = []
        curr_p_idx = start_p_idx

        for r_idx, row in enumerate(table.rows):
            cells_models: List[TableCellModel] = []
            for c_idx, cell in enumerate(row.cells):
                cell_paras: List[ParagraphModel] = []
                for doc_p in cell.paragraphs:
                    loc = SourceLocation(
                        location_type=SourceLocationType.TABLE,
                        paragraph_index=curr_p_idx,
                        table_index=table_idx,
                        row_index=r_idx,
                        cell_index=c_idx,
                    )
                    p_model = self._parse_paragraph(doc_p, curr_p_idx, loc)
                    cell_paras.append(p_model)
                    curr_p_idx += 1

                    p_links = self._extract_paragraph_hyperlinks(doc_p, p_model, relationships)
                    hyperlinks.extend(p_links)

                cells_models.append(
                    TableCellModel(
                        cell_index=c_idx,
                        row_index=r_idx,
                        table_index=table_idx,
                        paragraphs=cell_paras,
                    )
                )
            rows_models.append(
                TableRowModel(row_index=r_idx, table_index=table_idx, cells=cells_models)
            )

        return TableModel(table_index=table_idx, rows=rows_models)

    def _parse_header_footer(
        self,
        paragraphs: List[docx.text.paragraph.Paragraph],
        kind_name: str,
        start_p_idx: int,
        relationships: Dict[str, RelationshipModel],
        hyperlinks: List[HyperlinkModel],
    ) -> HeaderFooterModel:
        """Parse header or footer paragraphs into HeaderFooterModel."""
        hdr_ftr_paras: List[ParagraphModel] = []
        curr_p_idx = start_p_idx

        loc_type = (
            SourceLocationType.HEADER
            if "header" in kind_name
            else SourceLocationType.FOOTER
        )

        for doc_p in paragraphs:
            loc = SourceLocation(
                location_type=loc_type,
                paragraph_index=curr_p_idx,
                header_kind=kind_name if loc_type == SourceLocationType.HEADER else None,
                footer_kind=kind_name if loc_type == SourceLocationType.FOOTER else None,
            )
            p_model = self._parse_paragraph(doc_p, curr_p_idx, loc)
            hdr_ftr_paras.append(p_model)
            curr_p_idx += 1

            p_links = self._extract_paragraph_hyperlinks(doc_p, p_model, relationships)
            hyperlinks.extend(p_links)

        return HeaderFooterModel(kind=kind_name, paragraphs=hdr_ftr_paras)
